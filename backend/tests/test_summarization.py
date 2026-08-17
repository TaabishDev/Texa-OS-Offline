import os
import sys
import unittest
import tempfile
import asyncio
from pathlib import Path
from unittest.mock import AsyncMock, patch

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from agents.summarization_agent import SummarizationAgent
from docx import Document

class TestSummarization(unittest.IsolatedAsyncioTestCase):

    def setUp(self):
        self.agent = SummarizationAgent()
        self.temp_dir = tempfile.TemporaryDirectory()
        self.agent.temp_dir = Path(self.temp_dir.name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_file_validation(self):
        # Valid files
        self.assertTrue(self.agent.validate_file("test.pdf", 1000))
        self.assertTrue(self.agent.validate_file("paper.docx", 2000))
        
        # Invalid extensions
        with self.assertRaises(ValueError):
            self.agent.validate_file("image.png", 1000)
            
        # Exceeds size limit
        with self.assertRaises(ValueError):
            self.agent.validate_file("large.pdf", 30 * 1024 * 1024)

    def test_filename_sanitization(self):
        self.assertEqual(self.agent.sanitize_filename("../../etc/passwd.pdf"), "passwd.pdf")
        self.assertEqual(self.agent.sanitize_filename("valid_name-1.2.docx"), "valid_name-1.2.docx")
        self.assertEqual(self.agent.sanitize_filename("spaces in name.pdf"), "spaces_in_name.pdf")

    def test_docx_text_extraction(self):
        # Create a mock docx file
        doc_path = Path(self.temp_dir.name) / "test_doc.docx"
        doc = Document()
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph("This is a paragraph about S&T document analysis.")
        doc.save(str(doc_path))

        elements = self.agent.extract_text(doc_path)
        self.assertTrue(len(elements) >= 2)
        self.assertEqual(elements[0]["type"], "heading")
        self.assertEqual(elements[0]["text"], "Abstract")
        self.assertEqual(elements[1]["type"], "paragraph")
        self.assertEqual(elements[1]["text"], "This is a paragraph about S&T document analysis.")

    def test_section_detection(self):
        elements = [
            {"type": "heading", "text": "Introduction to AI", "page": 1},
            {"type": "paragraph", "text": "AI has evolved fast.", "page": 1},
            {"type": "heading", "text": "Methodology", "page": 2},
            {"type": "paragraph", "text": "We used TF-IDF overlap.", "page": 2}
        ]
        processed = self.agent.detect_sections(elements)
        self.assertEqual(processed[0]["section"], "INTRODUCTION")
        self.assertEqual(processed[1]["section"], "INTRODUCTION")
        self.assertEqual(processed[2]["section"], "METHODOLOGY")
        self.assertEqual(processed[3]["section"], "METHODOLOGY")

    def test_chunking_rules(self):
        elements = [
            {"type": "heading", "text": "Introduction", "page": 1, "section": "INTRODUCTION"},
            {"type": "paragraph", "text": "A" * 3000, "page": 1, "section": "INTRODUCTION"},
            {"type": "paragraph", "text": "B" * 4000, "page": 2, "section": "INTRODUCTION"},
        ]
        # Chunk size 5000, overlap 100
        chunks = self.agent.chunk_document(elements, chunk_size=5000, overlap=100)
        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0]["chunk_id"], 1)
        self.assertEqual(chunks[0]["page_start"], 1)
        self.assertEqual(chunks[1]["chunk_id"], 2)
        self.assertEqual(chunks[1]["page_start"], 2)

    @patch("agents.summarization_agent.SummarizationAgent.call_llm", new_callable=AsyncMock)
    async def test_map_reduce_pipeline(self, mock_call_llm):
        mock_call_llm.side_effect = [
            "Summary of chunk 1",
            "Summary of chunk 2",
            '{"executive_summary": "Unified document overview", "abstract": null, "objectives": "Verify S&T summarizer output", "methodology": "Mock", "key_findings": "Mock results", "important_results": null, "technologies_methods": null, "limitations": null, "future_work": null, "conclusion": "Approved", "key_terms": ["Test", "Mock"]}'
        ]

        document_id = "test-doc-id"
        self.agent.documents_db[document_id] = {
            "status": "processing",
            "progress": "",
            "chunks": [],
            "summary": {},
            "file_path": "",
            "filename": "test.pdf"
        }
        chunks = [
            {"chunk_id": 1, "section": "ABSTRACT", "page_start": 1, "page_end": 1, "text": "Chunk 1 content"},
            {"chunk_id": 2, "section": "RESULTS", "page_start": 2, "page_end": 2, "text": "Chunk 2 content"}
        ]
        
        # Test summarize_chunks
        semaphore = asyncio.Semaphore(3)
        summaries = await self.agent.summarize_chunks(document_id, chunks, semaphore)
        self.assertEqual(len(summaries), 2)
        self.assertEqual(summaries[0], "Summary of chunk 1")

        # Test reduce_summaries
        final_result = await self.agent.reduce_summaries(summaries)
        self.assertEqual(final_result["executive_summary"], "Unified document overview")
        self.assertEqual(final_result["conclusion"], "Approved")

    @patch("agents.summarization_agent.SummarizationAgent.call_llm", new_callable=AsyncMock)
    async def test_grounded_qa(self, mock_call_llm):
        mock_call_llm.return_value = "The document uses python-docx."
        doc_id = "test-qa-id"
        self.agent.documents_db[doc_id] = {
            "status": "completed",
            "progress": "Completed",
            "chunks": [
                {"chunk_id": 1, "section": "TECH", "page_start": 1, "page_end": 1, "text": "We implement python-docx parsing."}
            ],
            "summary": {},
            "file_path": ""
        }
        res = await self.agent.query_document(doc_id, "Which library is used?")
        self.assertEqual(res["answer"], "The document uses python-docx.")
        self.assertEqual(res["source_section"], "TECH")
        self.assertEqual(res["page_number"], "1-1")

if __name__ == "__main__":
    unittest.main()
